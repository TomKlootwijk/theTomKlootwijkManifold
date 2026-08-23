from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Tom\Documents\theTomKlootwijkManifold")
PACKAGE = ROOT / "UGTS_KC_Tom_Klootwijk_Signature_3_9_1_Package" / "UGTS_KC_Tom_Klootwijk_Signature_3_9_1_Package"
OUTPUT = PACKAGE / "report" / "UGTS_KC_3_9_1_Tom_Signature_Native_Android_ELI5_Technical_Report.docx"
SCREENSHOT = Path(r"C:\Users\Tom\AppData\Local\Temp\codex-ugts-deploy-check\ugts-deploy-check-fixed.png")


# compact_reference_guide preset tokens
FONT = "Calibri"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
CYAN = "008FA8"  # named cover/screenshot accent override
MUTED = "667085"
LIGHT = "F4F6F9"
TABLE_FILL = "E8EEF5"
GRID = "CCD5E0"
WHITE = "FFFFFF"
INK = "172033"
CAUTION = "7A5A00"
RISK = "9B1C1C"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run(run, *, size=None, bold=None, italic=None, color=INK, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def set_spacing(paragraph, *, before=0, after=6, line=1.25, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcMar"))
    if old is not None:
        tc_pr.remove(old)
    tc_mar = OxmlElement("w:tcMar")
    for side, value in CELL_MARGINS.items():
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_DXA
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:color"), GRID)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def style_cell(cell, *, bold=False, color=INK, size=9.3, align=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        set_spacing(paragraph, before=0, after=0, line=1.15)
        for run in paragraph.runs:
            set_run(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, *, aligns=None, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for idx, header in enumerate(headers):
        hdr.cells[idx].text = str(header)
        set_cell_shading(hdr.cells[idx], TABLE_FILL)
        style_cell(hdr.cells[idx], bold=True, color=NAVY, size=font_size, align=(aligns[idx] if aligns else WD_ALIGN_PARAGRAPH.LEFT))
    set_repeat_table_header(hdr)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = str(value)
            style_cell(row.cells[idx], size=font_size, align=(aligns[idx] if aligns else WD_ALIGN_PARAGRAPH.LEFT))
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    set_spacing(after, before=0, after=4, line=1.0)
    return table


def add_callout(doc, label, text, *, fill=LIGHT, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    set_spacing(paragraph, before=0, after=0, line=1.2)
    lead = paragraph.add_run(label + " ")
    set_run(lead, bold=True, color=label_color, size=10.5)
    body = paragraph.add_run(text)
    set_run(body, color=INK, size=10.5)
    set_table_geometry(table, [CONTENT_DXA])
    spacer = doc.add_paragraph()
    set_spacing(spacer, before=0, after=4, line=1.0)
    return table


def add_body(doc, text, *, bold_lead=None, italic=False, color=INK, after=6):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=after, line=1.25)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run(lead, bold=True, color=color)
        rest = p.add_run(text[len(bold_lead):])
        set_run(rest, italic=italic, color=color)
    else:
        run = p.add_run(text)
        set_run(run, italic=italic, color=color)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=3, after=8, line=1.0)
    run = p.add_run(text)
    set_run(run, size=9, italic=True, color=MUTED)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    if level == 1:
        if getattr(doc, "_seen_h1", False):
            p.paragraph_format.page_break_before = True
        else:
            doc._seen_h1 = True
    return p


def add_bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style="Normal")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(doc._bullet_num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
    set_spacing(p, before=0, after=4, line=1.25)
    set_run(p.add_run(text), color=INK)
    return p


def add_numbered(doc, title, explanation):
    p = doc.add_paragraph(style="Normal")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(doc._decimal_num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
    set_spacing(p, before=0, after=5, line=1.25)
    set_run(p.add_run(title + ": "), bold=True, color=NAVY)
    set_run(p.add_run(explanation), color=INK)
    return p


def add_numbering_instance(doc, abstract_id, *, restart=False):
    numbering = doc.part.numbering_part.element
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    if restart:
        level_override = OxmlElement("w:lvlOverride")
        level_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        level_override.append(start_override)
        num.append(level_override)
    numbering.append(num)
    return num_id


def add_numbering_definition(doc, *, bullet):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), FONT)
        r_fonts.set(qn("w:hAnsi"), FONT)
        r_pr.append(r_fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)
    return abstract_id, add_numbering_instance(doc, abstract_id)


def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, text, end])
    set_run(run, size=9, color=MUTED)


def add_source_path(doc, label, path):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=4, line=1.15)
    set_run(p.add_run(label + ": "), bold=True, size=9.2, color=NAVY)
    set_run(p.add_run(path), size=8.7, color=MUTED, font="Consolas")
    return p


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in (
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    doc._bullet_abstract_id, doc._bullet_num_id = add_numbering_definition(doc, bullet=True)
    doc._decimal_abstract_id, doc._decimal_num_id = add_numbering_definition(doc, bullet=False)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(hp, before=0, after=0, line=1.0)
    set_run(hp.add_run("UGTS-KC 3.9.1  |  Native Android Technical Explainer"), size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(fp, before=0, after=0, line=1.0)
    set_run(fp.add_run("Page "), size=9, color=MUTED)
    add_field(fp, "PAGE")
    set_run(fp.add_run(" of "), size=9, color=MUTED)
    add_field(fp, "NUMPAGES")


def add_cover(doc):
    spacer = doc.add_paragraph()
    set_spacing(spacer, before=64, after=0, line=1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=18, line=1.0)
    set_run(p.add_run("TOM KLOOTWIJK SIGNATURE EDITION"), size=10, bold=True, color=CYAN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=8, line=1.0)
    set_run(p.add_run("How a Tiny Substrate Becomes\na 120 Hz Native Android App"), size=28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=30, line=1.15)
    set_run(p.add_run("An Explain Like I'm 5 technical report for UGTS-KC 3.9.1"), size=14, color=DARK_BLUE)

    add_callout(
        doc,
        "The little-kid version:",
        "Tom's package is the recipe and the toy world. Android's standard tools are the factory. The Poco phone is the stage. The factory changes the recipe into phone-shaped machine code, but it does not invent the rules, scene, or visual identity.",
        fill="EEF8FA",
        label_color=CYAN,
    )

    rows = [
        ("Prepared for / requested author", "Tom Klootwijk"),
        ("Requester-supplied identifier", "NL200678942"),
        ("Requester-supplied date of birth", "10-07-1990"),
        ("Version examined", "UGTS-KC 3.9.1 - Tom Klootwijk Signature Edition"),
        ("Physical target", "Poco X7 Pro 12 GB, model 2412DPC0AG"),
        ("Measurement date", "23 August 2026"),
    ]
    add_table(doc, ["Record", "Value"], rows, [2700, 6660], font_size=9.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=16, after=0, line=1.1)
    set_run(
        p.add_run("Attribution and identity details are requester supplied. This technical report is not an identity check, legal ownership ruling, or cryptographic personal signature."),
        size=8.8,
        italic=True,
        color=MUTED,
    )
    doc.add_page_break()


def build_report():
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "UGTS-KC 3.9.1 Tom Signature Native Android ELI5 Technical Report"
    props.subject = "Source lineage, native Android architecture, calculated workload, and Poco X7 Pro measurements"
    props.author = "Prepared for Tom Klootwijk"
    props.last_modified_by = "OpenAI Codex"
    props.keywords = "UGTS-KC, Tom Klootwijk, NativeActivity, Android, OpenGL ES, Poco X7 Pro, performance"
    props.comments = "Based on requester-supplied source package and on-device measurements captured 23 August 2026."

    add_cover(doc)

    add_heading(doc, "1. The answer in one minute", 1)
    add_callout(
        doc,
        "Bottom line:",
        "The deployed app is a small native C++ program whose game rules, scene, quality policy, and compact KC3D391 data come from the supplied 3.9.1 package. Existing Android tools compile and host it. No commercial game engine and no network service is involved.",
    )
    add_body(doc, "Think of the project as five simple objects:")
    add_bullet(doc, "The blueprint: project.json describes the arena, meshes, materials, rules, camera, light, and device tiers.")
    add_bullet(doc, "The lunchbox: signature_scene.kc3d packs the blueprint into 21,735 bytes for fast native loading.")
    add_bullet(doc, "The brain: 12 C++ source/header files implement the loop, input, movement, loading, device choice, adaptive quality, and drawing.")
    add_bullet(doc, "The painter: OpenGL ES 3.0 asks the Mali-G720 GPU to draw colored triangles.")
    add_bullet(doc, "The stage manager: Android NativeActivity supplies the window, lifecycle, touch/gamepad events, and access to the GPU driver.")

    rows = [
        ("Presented frame rate", "120.07 FPS mean", "374 SurfaceFlinger presentation intervals"),
        ("Frame interval", "8.329 ms mean", "8.528 ms p95; 8.770 ms maximum"),
        ("App CPU", "61.92% of one core", "7.74% of total 8-core capacity"),
        ("Memory", "136.18 MiB PSS", "251.48 MiB RSS; 64.49 MiB graphics"),
        ("Thermal snapshot", "Status 0 - no throttling", "CPU/GPU 41.7 C; skin/battery 32.4 C"),
    ]
    add_table(doc, ["Measured item", "Result", "Meaning"], rows, [2500, 2300, 4560], font_size=9.1)

    if SCREENSHOT.exists():
        p = doc.add_paragraph()
        p.paragraph_format.page_break_before = True
        p.paragraph_format.keep_with_next = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        shape = run.add_picture(str(SCREENSHOT), width=Inches(6.45))
        doc_pr = shape._inline.docPr
        doc_pr.set("descr", "UGTS-KC Tom Klootwijk Signature Arena 3D rendering on a Poco X7 Pro in landscape orientation")
        add_caption(doc, "Figure 1. The actual deployed native scene captured from the Poco X7 Pro after the event-loop correction.")

    add_callout(
        doc,
        "Important measurement boundary:",
        "These are short, foreground, USB-connected debug-build measurements with a static camera. They prove that this build presents smoothly at 120 Hz in this test. They are not a 30-minute heat, battery, gameplay, or release-build certification.",
        fill="FFF8E6",
        label_color=CAUTION,
    )
    add_heading(doc, "What this picture confirms", 2)
    add_bullet(doc, "The native activity reached a real EGL/OpenGL ES surface instead of remaining on a black window.")
    add_bullet(doc, "The KC3D scene, shaders, camera, colors, depth, and shared meshes were loaded and drawn on the Poco.")
    add_bullet(doc, "A screenshot proves visible output; the timing measurements later in this report are the separate performance evidence.")

    add_heading(doc, "2. Where the deployed app came from", 1)
    add_body(
        doc,
        "The source package already contains the authoring model, compiler, compact scene, Android template, checked-in Android project, validation records, and attribution files. The standalone project that was built is a copy of that Android tree.",
    )
    add_heading(doc, "A file-by-file lineage check", 2)
    add_body(doc, "Twenty-seven build/runtime files were compared by SHA-256 between the supplied package and the deployed project:")
    add_bullet(doc, "26 files are byte-for-byte identical.")
    add_bullet(doc, "1 file differs: app/src/main/cpp/main.cpp.")
    add_bullet(doc, "The difference is a minimal Android event-loop correction. It processes one looper result, returns to the outer loop, and lets the first frame render after focus changes.")
    add_bullet(doc, "No scene, mesh, material, shader, gameplay rule, device profile, or art asset was added by that correction.")

    add_callout(
        doc,
        "Why this honesty matters:",
        "The first archive version compiled and launched, but its stale infinite-wait timeout left the display black. The deployed APK therefore is not byte-for-byte the untouched archive. It is the archive plus this small platform-integration fix.",
        fill="FDECEC",
        label_color=RISK,
    )

    rows = [
        ("Canonical project-content hash", "e86bebe946cb6a4f7e418143869191c4900bed22d2773054e0b1bf0a6493c244"),
        ("KC3D391 scene-pack SHA-256", "485bb8a5ed7d4632577195dc0a889b4e6d3ea13ebc55482c5ea40761729ba437"),
        ("Deployed debug APK SHA-256", "e3229b25529d222dd1ff5d0f98ac5ddf2fb283d73b2a6b4208a4b6df054821cf"),
        ("Application ID", "nl.tomklootwijk.ugtskc.signature.pocox7pro"),
        ("Version", "391 / 3.9.1-poco-x7-pro"),
    ]
    add_table(doc, ["Integrity record", "Value"], rows, [3000, 6360], font_size=8.7)

    add_heading(doc, "Authorship wording used in this report", 2)
    add_body(
        doc,
        "This report describes the package as Tom Klootwijk's requested original/signature work because that is the attribution supplied by the requester and repeated in the package. The package itself says the identity, authorship, and rights assertion was not independently verified. The identifier NL200678942 and date 10-07-1990 are therefore presented as requester-supplied record data, not as a government verification.",
    )
    add_heading(doc, "3. From tiny description to running phone", 1)
    add_body(doc, "Here is the whole factory line in child-friendly language, followed by the grown-up technical meaning.")
    add_numbered(doc, "Write the toy-world recipe", "The editable project.json names every object, color, transform, rule, device tier, camera, and light. This is the source of truth.")
    add_numbered(doc, "Fold the recipe into a lunchbox", "The package compiler converts the JSON into the little-endian KC3D391 binary scene pack. It stores validated counts, strings, meshes, materials, nodes, quality tiers, and target profiles.")
    add_numbered(doc, "Read the factory card", "Gradle 8.13 and Android Gradle Plugin 8.13.2 read the application ID, SDK levels, ARM64 Poco flavor, assets, and native-build instructions.")
    add_numbered(doc, "Build the metal brain", "CMake 3.22.1 and Ninja ask Android NDK r29's Clang compiler to turn six C++20 implementation files into libugts_kc_native.so for arm64-v8a.")
    add_numbered(doc, "Pack the backpack", "AAPT2 and the Android packaging tasks combine the manifest, tiny resources, native library, KC3D scene, source JSON, inspection JSON, and two shader files into one APK.")
    add_numbered(doc, "Put it on the phone", "ADB streams the APK to the Poco. Android verifies the debug signature and installs the application under its package name.")
    add_numbered(doc, "Open the stage", "Android starts android.app.NativeActivity, creates a native window, and loads libugts_kc_native.so. No handwritten Java or Kotlin activity is required.")
    add_numbered(doc, "Unpack the scene", "The C++ asset loader opens signature_scene.kc3d, validates its format, uploads four meshes to GPU buffers, and compiles the two GLSL ES shaders.")
    add_numbered(doc, "Choose the Poco setting", "The explicit build flavor selects poco_x7_pro_12gb, requests 120 FPS, and starts with signature_ultra at full render scale.")
    add_numbered(doc, "Play and paint", "The loop reads input, advances fixed-step gameplay, draws the scene into an offscreen framebuffer, copies it to the phone window, and swaps on vertical sync.")

    add_callout(
        doc,
        "What Android tooling contributes:",
        "It contributes the compiler, packager, operating-system window, input plumbing, system libraries, graphics API, GPU driver, installer, and debug signature. It does not supply the scene design, KC3D data model, node rules, quality policy, or game-specific C++ logic.",
    )
    add_heading(doc, "4. What is actually used", 1)
    rows = [
        ("Authoring source", "Build-time", "project.json is the editable record. It is compiled into KC3D391."),
        ("KC3D391 scene", "Runtime", "The engine reads 4 meshes, 8 materials, 66 nodes, 5 quality tiers, and 4 target profiles."),
        ("C++20 engine", "Runtime", "Lifecycle, input, fixed-step movement, collision checks, camera, quality logic, loading, and rendering."),
        ("NativeActivity + native_app_glue", "Runtime/platform", "Android lifecycle, native window, input queue, and looper integration."),
        ("EGL + OpenGL ES 3.0", "Runtime/platform", "Creates the graphics context, buffers, shaders, draws, blits, and presents frames."),
        ("Mali-G720 MC7 driver", "Runtime/device", "Executes the GLES commands on the Poco GPU."),
        ("JNI PowerManager call", "Runtime/platform", "Reads Android's current thermal status."),
        ("Android build toolchain", "Build-time only", "Gradle, AGP, JDK, CMake, Ninja, and NDK turn source into the ARM64 APK; they do not run as the game."),
    ]
    add_table(doc, ["Piece", "When used", "Actual job"], rows, [2450, 1700, 5210], font_size=9.0)

    add_heading(doc, "Things shipped but not opened by the current game loop", 2)
    add_bullet(doc, "assets/project.json is kept in the APK for inspection/provenance, but the C++ runtime opens signature_scene.kc3d instead.")
    add_bullet(doc, "assets/scene-pack-inspection.json is also packaged for human/tool inspection, but the runtime does not read it.")

    add_heading(doc, "Things in the wider package that are not running on this phone", 2)
    add_bullet(doc, "The Python reference oracle, schema validators, CLI, tests, 2D browser runtime, HTML5 exporter, glTF/USDA tools, reports, and diagrams.")
    add_bullet(doc, "Unity, Unreal Engine, AndroidX, a web view, a Java game engine, cloud services, analytics, advertising, or network assets.")
    add_bullet(doc, "Vulkan: it is only declared optional and documented as a future hook. The renderer is GLES3.")
    add_bullet(doc, "4D execution: it remains a design-only TODO and is not part of the running binary.")
    add_bullet(doc, "Audio and textures: the current APK contains neither. The visual scene uses primitive meshes and flat/emissive material colors.")
    add_bullet(doc, "Android permissions: the manifest requests none.")
    add_heading(doc, "5. What happens inside one frame", 1)
    doc._decimal_num_id = add_numbering_instance(doc, doc._decimal_abstract_id, restart=True)
    add_numbered(doc, "Listen", "Poll one Android looper result. Touch, keyboard, gamepad, focus, and window events enter here.")
    add_numbered(doc, "Measure time", "Compute elapsed time and clamp it to 0.1 second so a long pause cannot create an enormous physics jump.")
    add_numbered(doc, "Tick the rules", "Advance the world in exact 1/120-second slices: gravity, movement, floor/bounds, rotations, and player-vs-node contacts.")
    add_numbered(doc, "Move the camera", "Orbit around the player target using yaw, pitch, and pinch distance.")
    add_numbered(doc, "Check the thermostat", "About once each second, compare measured FPS and Android thermal status with the adaptive-quality thresholds.")
    add_numbered(doc, "Paint offscreen", "Clear color/depth, set global light/camera uniforms, then issue one indexed draw for each visible live node.")
    add_numbered(doc, "Copy and show", "Blit the offscreen color image to the Android window and call eglSwapBuffers with swap interval 1.")

    add_heading(doc, "Scene workload", 2)
    rows = [
        ("Cube", "45", "12", "540"),
        ("Floor", "1", "2", "2"),
        ("Pyramid", "6", "6", "36"),
        ("Sphere", "14", "440", "6,160"),
        ("Total", "66", "-", "6,738"),
    ]
    add_table(
        doc,
        ["Shared mesh", "Node draws", "Triangles per draw", "Triangles per frame"],
        rows,
        [2200, 1800, 2380, 2980],
        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=9.2,
    )
    add_body(doc, "The four meshes contain only 306 unique vertex records and 460 unique triangles. Reusing them across 66 nodes produces the larger per-frame workload without duplicating the mesh buffers.")
    add_callout(
        doc,
        "At 120 presented frames per second:",
        "66 draws/frame becomes 7,920 draw calls/second; 6,738 triangles/frame becomes 808,560 triangles/second; 20,214 submitted indices/frame becomes 2,425,680 indices/second. This is a modest geometry load for a Mali-G720. The current renderer favors simplicity over batching: it does not use GPU instancing.",
    )
    add_heading(doc, "6. Pixel work and memory, calculated", 1)
    add_body(doc, "The app is rotated to 2712 x 1220 landscape pixels. At full scale:")
    add_bullet(doc, "Pixels per rendered image = 2712 x 1220 = 3,308,640 pixels (3.309 megapixels).")
    add_bullet(doc, "Offscreen color = 4 bytes/pixel; depth is budgeted as 4 bytes/pixel for the DEPTH_COMPONENT24 renderbuffer allocation.")
    add_bullet(doc, "Offscreen color + depth = 3,308,640 x 8 = 26,469,120 bytes = 25.24 MiB.")
    add_bullet(doc, "The renderer clears the offscreen target and blits a full color image to the window: a minimum of two full-screen pixel passes before counting shaded geometry or system composition.")
    add_bullet(doc, "At 120 FPS, that clear-plus-blit floor is about 794.1 million pixel operations per second.")

    rows = [
        ("Signature ultra", "1.00", "2712 x 1220", "3.309", "100.0%", "25.24", "794.1"),
        ("High", "0.92", "2495 x 1122", "2.799", "84.6%", "21.36", "671.8"),
        ("Balanced", "0.82", "2224 x 1000", "2.224", "67.2%", "16.97", "533.8"),
        ("Safe", "0.68", "1844 x 830", "1.531", "46.3%", "11.68", "367.3"),
        ("Thermal", "0.55", "1492 x 671", "1.001", "30.3%", "7.64", "240.3"),
    ]
    add_table(
        doc,
        ["Tier", "Scale", "Internal size", "Mpix", "% full", "FBO MiB", "2-pass Mpix/s at 120"],
        rows,
        [1850, 850, 1680, 900, 920, 1100, 2060],
        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=8.2,
    )
    add_body(doc, "The measured graphics allocation was 64.49 MiB, about 2.56 times the calculated offscreen color-plus-depth pair. The difference is reasonable because Android's window buffers, driver allocations, buffer queues, and other graphics bookkeeping are outside that simple two-buffer calculation.")
    add_callout(
        doc,
        "A subtle but important fact:",
        "All five max-visible-node limits are above this scene's 66 nodes; even the thermal cap is 160. Therefore the current scene does not reduce draw count when quality steps down. In practice, the implemented fallback mainly reduces pixel resolution.",
        fill="FFF8E6",
        label_color=CAUTION,
    )
    add_heading(doc, "7. What the Poco actually measured", 1)
    add_body(doc, "Test conditions: app in the foreground, static camera, screen awake, connected by USB/ADB, debug build, Android 16/API 36, landscape 2712 x 1220, active display mode 120 Hz.")
    rows = [
        ("Display mode", "120.00001 Hz", "Android display service; active mode ID 3"),
        ("Presentation sample", "374 intervals", "Three cleared SurfaceFlinger windows"),
        ("Mean presentation interval", "8.329 ms", "Equivalent to 120.07 presented FPS"),
        ("Median / p95 / p99", "8.324 / 8.528 / 8.640 ms", "Tight pacing in the sampled static view"),
        ("Maximum interval", "8.770 ms", "No interval exceeded 12.5 ms"),
        ("Intervals above 1.5 vsync", "0 of 374", "0.00% in this short sample"),
        ("Process CPU", "61.92% of one core", "Mean of three 3-second /proc samples"),
        ("Whole-chip equivalent", "7.74% of 8-core capacity", "Simple 61.92 / 8 normalization"),
        ("Total PSS", "139,447 KiB = 136.18 MiB", "Proportional memory snapshot"),
        ("Total RSS", "257,518 KiB = 251.48 MiB", "Resident pages, including shared mappings"),
        ("Graphics", "66,042 KiB = 64.49 MiB", "EGL + GL memory reported by dumpsys"),
        ("Thermal status", "0", "Android says no thermal throttling state"),
        ("HAL temperature snapshot", "CPU/GPU 41.7 C", "Battery/skin 32.4 C after restart/sample"),
        ("Start after force-stop", "235.3 ms mean", "225-247 ms across three cached launches"),
    ]
    add_table(doc, ["Metric", "Measured result", "Method / meaning"], rows, [2600, 2600, 4160], font_size=8.6)

    add_heading(doc, "How to read these numbers", 2)
    add_bullet(doc, "SurfaceFlinger measures when frames were presented. It is not a direct GPU render-duration trace.")
    add_bullet(doc, "The 120.07 value is timing arithmetic around a nominal 120 Hz mode; report it as approximately 120 FPS, not as an overclock.")
    add_bullet(doc, "CPU percent is process time. Some display/GPU work occurs in Android system or kernel processes and is not fully charged to the app process.")
    add_bullet(doc, "RSS counts shared mappings more generously than PSS. PSS is the better single number for the app's proportional memory cost.")
    add_bullet(doc, "The 235.3 ms result is ActivityManager's start-completion time after force-stop with warm file caches, not a camera-recorded first-interactive-pixel measurement.")
    add_bullet(doc, "No battery-current, wattage, GPU counter, or long-duration thermal test was captured.")
    add_heading(doc, "8. The quality thermostat - and its current limits", 1)
    add_body(doc, "The quality controller is like a parent turning down a toy's detail when the phone gets tired.")
    add_bullet(doc, "Stress is declared when Android thermal status is 3 or higher, or measured FPS is below 82% of 120 = 98.4 FPS.")
    add_bullet(doc, "The configured stress timer is 1.5 seconds. Because the engine evaluates it only about once per second, a continuous problem normally needs two stressed checks before one tier is dropped.")
    add_bullet(doc, "Recovery requires thermal status 1 or lower and FPS at or above 96% of 120 = 115.2 FPS for 8 seconds before moving up one tier.")
    add_bullet(doc, "Only one tier changes per decision, preventing an immediate jump from ultra to thermal or back.")

    rows = [
        ("0", "Signature ultra", "1.00", "1024", "120"),
        ("1", "High", "0.92", "720", "90"),
        ("2", "Balanced", "0.82", "480", "60"),
        ("3", "Safe", "0.68", "280", "60"),
        ("4", "Thermal", "0.55", "160", "45"),
    ]
    add_table(doc, ["Index", "Tier", "Render scale", "Node cap", "Stored target FPS"], rows, [900, 2480, 1800, 1800, 2380], font_size=9.0)

    add_heading(doc, "Critical implementation facts", 2)
    add_callout(
        doc,
        "1. Stored target FPS values do not retune presentation:",
        "After startup, the app continues requesting 120 FPS and the adaptive comparison continues using the profile's 120 FPS target. The 90/60/45 values stored in lower tiers are parsed but are not used to change the frame-rate request.",
        fill="FFF8E6",
        label_color=CAUTION,
    )
    add_callout(
        doc,
        "2. Three quality fields are presently descriptive only:",
        "MSAA sample count, post-processing, and shadow quality are loaded from the scene pack, but the current GLES3 renderer does not implement them. There is no MSAA attachment, post-processing pass, or shadow pass.",
        fill="FFF8E6",
        label_color=CAUTION,
    )
    add_callout(
        doc,
        "3. Node caps do not affect this 66-node demo:",
        "Every tier permits at least 160 visible nodes. There is also no frustum culling or distance sorting; nodes are considered in stored order until the cap is reached.",
        fill="FFF8E6",
        label_color=CAUTION,
    )
    add_body(doc, "These are not reasons to dismiss the design. They are the exact boundary between the compact policy encoded by the substrate and the subset that the present renderer actively enforces.")
    add_heading(doc, "9. How compact is it?", 1)
    rows = [
        ("Editable project JSON", "160,201 bytes", "Human-readable source record"),
        ("KC3D391 binary pack", "21,735 bytes", "86.43% smaller; JSON is 7.37x larger"),
        ("C++ implementation", "12 files / 1,081 lines", "42,874 bytes; 1,000 nonblank lines"),
        ("GLSL shader source", "1,093 bytes", "One vertex shader + one fragment shader"),
        ("Debug APK", "1,173,398 bytes", "1.119 MiB total archive"),
        ("ARM64 native library", "1,121,752 bytes", "Uncompressed in APK; about 95.6% of APK size"),
        ("Tiny DEX", "1,396 bytes", "No handwritten Java/Kotlin game code"),
    ]
    add_table(doc, ["Item", "Footprint", "Interpretation"], rows, [2650, 2350, 4360], font_size=9.0)

    add_heading(doc, "Why the native library is the big piece", 2)
    add_body(doc, "The project requests the static C++ standard library (c++_static). That library support is folded into libugts_kc_native.so instead of being a separate runtime dependency. The stripped binary's text section is about 1.081 MB, so the native code/support dominates the APK while the custom scene and shader data remain tiny.")

    add_heading(doc, "Tiny substrate does not mean zero platform", 2)
    add_body(doc, "The APK is small because it reuses facilities already present on Android: the operating system, Linux process model, NativeActivity framework, EGL, OpenGL ES API, Mali driver, display compositor, input system, power service, and installer. Those platform components are substantial, but they are not copied into the APK and they do not replace the package's game-specific design.")
    add_callout(
        doc,
        "Most accurate one-sentence description:",
        "A compact Tom Klootwijk-attributed data/engine package is cross-compiled against standard Android native interfaces, then the Poco's existing OS and GPU execute that package directly.",
    )
    add_heading(doc, "10. What this proves, and what it does not", 1)
    add_heading(doc, "Demonstrated in this session", 2)
    add_bullet(doc, "The Poco-specific ARM64 debug APK compiles with SDK 36, NDK r29, CMake 3.22.1, and AGP/Gradle 8.13.x.")
    add_bullet(doc, "It installs, launches, stays in the foreground, loads the Mali-G720 GLES3 driver, selects the 12 GB Poco profile, and visibly renders the scene.")
    add_bullet(doc, "The active Android display mode is 120 Hz and the sampled presented-frame cadence is approximately 120 FPS with tight short-window pacing.")
    add_bullet(doc, "Short CPU, memory, launch, and thermal snapshots were captured and are reported with their methods.")
    add_bullet(doc, "The deployed source lineage matches the supplied Android project except for the disclosed event-loop correction.")

    add_heading(doc, "Not demonstrated", 2)
    add_bullet(doc, "Sustained 120 FPS during 30-60 minutes of active gameplay, charging, high ambient temperature, or low-battery conditions.")
    add_bullet(doc, "Battery drain, watts, GPU utilization/counters, per-pass GPU duration, or energy efficiency.")
    add_bullet(doc, "Release-mode performance, release signing, Play Store readiness, security review, or production certification.")
    add_bullet(doc, "Robust general physics, skeletal animation, networking, multiplayer, anti-cheat, Vulkan rendering, or a 4D runtime.")
    add_bullet(doc, "Independent verification of identity, authorship, ownership, or third-party rights.")

    add_heading(doc, "Most useful next benchmark", 2)
    add_body(doc, "For a production performance claim, the next honest test is a 30-minute release-build run with automated movement/camera input, Perfetto CPU/GPU counters where available, frame timeline data, battery current, and temperature samples every minute. That would test sustained behavior instead of the lightweight static view measured here.")
    add_heading(doc, "11. Evidence record and formulas", 1)
    add_heading(doc, "Primary package records", 2)
    add_source_path(doc, "Package root", str(PACKAGE))
    add_source_path(doc, "Signature record", r"signature\TOM_KLOOTWIJK_SIGNATURE_EDITION.json")
    add_source_path(doc, "Native contract", r"spec\MOBILE_3D_AND_ANDROID_CONTRACT.md")
    add_source_path(doc, "Scene-pack format", r"spec\NATIVE_SCENE_PACK_3_9_1.md")
    add_source_path(doc, "Evidence boundary", r"docs\EVIDENCE_BOUNDARY.md")
    add_source_path(doc, "Android guide", r"docs\ANDROID_NATIVE_GUIDE.md")
    add_source_path(doc, "Checked-in project", r"android\UGTSKC391Signature")
    add_source_path(doc, "Editable scene", r"examples\tom_signature_arena_3d\project.json")

    add_heading(doc, "Deployed implementation records", 2)
    add_source_path(doc, "Built project", str(ROOT / "UGTS_KC_3_9_1_Tom_Signature_Android_Source" / "UGTSKC391Signature"))
    add_source_path(doc, "Event-loop correction", r"app\src\main\cpp\main.cpp")
    add_source_path(doc, "Runtime engine", r"app\src\main\cpp\engine.cpp")
    add_source_path(doc, "GLES3 renderer", r"app\src\main\cpp\renderer_gles3.cpp")
    add_source_path(doc, "Compact scene", r"app\src\main\assets\signature_scene.kc3d")

    add_heading(doc, "Calculation formulas", 2)
    add_bullet(doc, "Triangles/frame = sum(node count using mesh x mesh index count / 3) = 6,738.")
    add_bullet(doc, "Presented FPS = 1000 / mean presentation interval in milliseconds = 1000 / 8.329 = 120.07.")
    add_bullet(doc, "Full pixels/frame = 2712 x 1220 = 3,308,640.")
    add_bullet(doc, "Offscreen color+depth MiB = width x height x (4+4) / 1,048,576 = 25.24 MiB.")
    add_bullet(doc, "Minimum two-pass pixel rate = pixels x 2 x 120 = 794,073,600 pixel operations/second.")
    add_bullet(doc, "Process CPU = process CPU seconds / wall seconds x 100; whole-chip normalization = one-core percent / 8.")
    add_bullet(doc, "PSS MiB = dumpsys KiB / 1024; 139,447 / 1024 = 136.18 MiB.")
    add_bullet(doc, "Pack reduction = 1 - 21,735 / 160,201 = 86.43%.")

    add_heading(doc, "Measurement interfaces", 2)
    measurement_methods = add_body(doc, "Android SurfaceFlinger --latency supplied presentation timestamps; dumpsys display supplied the active refresh mode; /proc/<pid>/stat supplied process CPU time; dumpsys meminfo supplied PSS/RSS/graphics memory; dumpsys thermalservice supplied status and temperatures; am start -W supplied ActivityManager start timing.")
    measurement_methods.paragraph_format.keep_together = True

    add_heading(doc, "Small glossary", 2)
    add_bullet(doc, "Substrate: the compact rules, records, formats, and source that define the work.")
    add_bullet(doc, "APK: the installable Android package placed on the phone.")
    add_bullet(doc, "Native: machine code compiled for the phone's ARM64 processor instead of game logic running in JavaScript or a managed game engine.")
    add_bullet(doc, "Frame: one complete picture presented to the display.")
    add_bullet(doc, "PSS: the app's private memory plus its fair share of shared memory.")
    add_bullet(doc, "Vsync: the display's regular heartbeat; at 120 Hz one beat is about 8.333 ms.")

    add_callout(
        doc,
        "Final ELI5 summary:",
        "Tom's compact files tell the phone what world exists and how it behaves. Standard Android tools translate those files into the phone's language. Android opens a window, the C++ brain advances the toy world 120 times a second, and the Mali GPU colors the triangles. In the short test, the Poco kept pace with the 120 Hz screen without a sampled missed-vsync interval or thermal warning.",
        fill="EEF8FA",
        label_color=CYAN,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(str(OUTPUT))


if __name__ == "__main__":
    build_report()
